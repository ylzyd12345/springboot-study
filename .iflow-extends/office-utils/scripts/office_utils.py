#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
办公文档处理工具库 - 优化版 v2.0
输出标准化数据结构，供AI理解和处理各类Office任务

核心设计原则：
1. 输出结构标准化 - 所有函数返回统一格式的字典
2. 数据自描述 - 包含元信息，AI可直接理解
3. 渐进式读取 - 支持只读结构/读部分数据/读全部数据
4. 错误友好 - 出错时返回错误信息而非抛异常
5. 性能优先 - 智能检测空行，避免无效遍历

使用方式:
    from tools import *
    
    # 快速了解文件结构（不读数据）
    info = get_file_info('file.xlsx')
    
    # 读取数据（可指定范围）
    data = read_excel('file.xlsx', preview_rows=10)
    
    # 分析目录
    results = analyze_directory('.')

"""

import os
import re
import csv
import subprocess
import json
import logging
from pathlib import Path
from datetime import datetime
from collections import Counter
from typing import List, Dict, Any, Optional, Tuple, Union, Callable
from functools import lru_cache

# ============================================================================
# 配置常量
# ============================================================================

class Config:
    """全局配置"""
    # 性能相关
    MAX_SCAN_ROWS = 10000          # 最大扫描行数
    EMPTY_ROW_THRESHOLD = 100     # 连续空行阈值，超过则停止
    MAX_PREVIEW_COLS = 30         # 预览最大列数
    MAX_HEADER_ROWS = 5           # 最大表头行数
    TYPE_INFER_SAMPLE_SIZE = 50   # 类型推断采样行数
    
    # 缓存相关
    CACHE_ENABLED = True
    CACHE_MAX_SIZE = 32
    
    # 日志 - MCP 模式下禁用以避免干扰通信
    LOG_LEVEL = logging.WARNING


# 配置日志 - 输出到 stderr 避免干扰 MCP stdout 通信
logging.basicConfig(
    level=Config.LOG_LEVEL,
    format='%(asctime)s - %(levelname)s - %(message)s',
    stream=__import__('sys').stderr
)
logger = logging.getLogger(__name__)


# ============================================================================
# 预编译正则表达式（性能优化 #4）
# ============================================================================

RE_NUMERIC = re.compile(r'^-?[\d,.]+$')
RE_DATE = re.compile(r'\d{4}[-/]\d{1,2}[-/]\d{1,2}')
RE_PERCENTAGE = re.compile(r'^-?[\d.]+%$')
RE_CURRENCY = re.compile(r'^[¥$€£]?-?[\d,.]+[¥$€£]?$')


# ============================================================================
# 文件格式识别
# ============================================================================

# 支持的格式
EXCEL_NEW = {'.xlsx'}
EXCEL_OLD = {'.xls', '.et'}
WORD_NEW = {'.docx'}
WORD_OLD = {'.doc', '.wps'}
PPT_NEW = {'.pptx'}
PPT_OLD = {'.ppt', '.dps'}
PDF = {'.pdf'}
CSV_FORMATS = {'.csv', '.tsv'}

ALL_OFFICE = EXCEL_NEW | EXCEL_OLD | WORD_NEW | WORD_OLD | PPT_NEW | PPT_OLD | PDF | CSV_FORMATS


def get_file_type(filepath: str) -> Tuple[Optional[str], bool]:
    """
    识别文件类型
    
    Returns:
        (类型, 是否需要转换): ('excel', True) 表示是旧版Excel需要转换
    """
    suffix = Path(filepath).suffix.lower()
    
    if suffix in EXCEL_NEW:
        return 'excel', False
    elif suffix in EXCEL_OLD:
        return 'excel', True
    elif suffix in WORD_NEW:
        return 'word', False
    elif suffix in WORD_OLD:
        return 'word', True
    elif suffix in PPT_NEW:
        return 'ppt', False
    elif suffix in PPT_OLD:
        return 'ppt', True
    elif suffix in PDF:
        return 'pdf', False
    elif suffix in CSV_FORMATS:
        return 'csv', False
    return None, False


def scan_office_files(directory: str = '.', recursive: bool = False) -> List[Path]:
    """
    扫描目录下所有Office文件
    
    Args:
        directory: 目录路径
        recursive: 是否递归扫描子目录
    """
    files = []
    directory = Path(directory)
    
    if recursive:
        for f in directory.rglob('*'):
            if f.is_file() and f.suffix.lower() in ALL_OFFICE and not f.name.startswith('~'):
                files.append(f)
    else:
        for f in directory.iterdir():
            if f.is_file() and f.suffix.lower() in ALL_OFFICE and not f.name.startswith('~'):
                files.append(f)
    
    return sorted(files)


# ============================================================================
# 格式转换
# ============================================================================

def convert_to_new_format(filepath: str, output_dir: str = None) -> Optional[Path]:
    """
    将旧版Office文件转换为新版格式
    需要安装 LibreOffice: brew install --cask libreoffice
    
    Args:
        filepath: 源文件路径
        output_dir: 输出目录，默认同目录
    
    Returns:
        转换后的文件路径
    """
    filepath = Path(filepath)
    output_dir = output_dir or str(filepath.parent)
    
    file_type, needs_convert = get_file_type(str(filepath))
    if not needs_convert:
        return filepath
    
    # 确定目标格式
    format_map = {'excel': 'xlsx', 'word': 'docx', 'ppt': 'pptx'}
    target_format = format_map.get(file_type)
    if not target_format:
        return None
    
    # LibreOffice 路径
    soffice_paths = [
        '/Applications/LibreOffice.app/Contents/MacOS/soffice',  # macOS
        '/usr/bin/soffice',  # Linux
        'soffice',  # PATH中
    ]
    
    soffice = None
    for p in soffice_paths:
        if os.path.exists(p) or p == 'soffice':
            soffice = p
            break
    
    if not soffice:
        logger.warning(f"未找到LibreOffice，请手动转换: {filepath}")
        return None
    
    try:
        cmd = [soffice, '--headless', '--convert-to', target_format, 
               '--outdir', output_dir, str(filepath)]
        subprocess.run(cmd, check=True, capture_output=True, timeout=60)
        new_path = Path(output_dir) / f"{filepath.stem}.{target_format}"
        logger.info(f"已转换: {filepath.name} -> {new_path.name}")
        return new_path
    except subprocess.TimeoutExpired:
        logger.error(f"转换超时: {filepath.name}")
        return None
    except Exception as e:
        logger.error(f"转换失败: {filepath.name}, 错误: {e}")
        return None


# ============================================================================
# 标准化输出结构
# ============================================================================

def _make_response(success: bool, data: Any = None, error: str = None, 
                   meta: Dict = None) -> Dict[str, Any]:
    """创建标准化响应结构"""
    return {
        'success': success,
        'data': data,
        'error': error,
        'meta': meta or {},
        'timestamp': datetime.now().isoformat()
    }


# ============================================================================
# 工具函数
# ============================================================================

def _human_size(size_bytes: int) -> str:
    """转换为人类可读的文件大小"""
    for unit in ['B', 'KB', 'MB', 'GB']:
        if size_bytes < 1024:
            return f"{size_bytes:.1f}{unit}"
        size_bytes /= 1024
    return f"{size_bytes:.1f}TB"


def _is_row_empty(row: Union[list, tuple]) -> bool:
    """检查行是否为空"""
    if not row:
        return True
    return all(v is None or (isinstance(v, str) and not v.strip()) for v in row)


def _check_file_accessible(filepath: Path) -> Tuple[bool, str]:
    """
    检查文件是否可访问（优化 #10）
    
    Returns:
        (是否可访问, 错误信息)
    """
    if not filepath.exists():
        return False, f"文件不存在: {filepath}"
    
    try:
        # 尝试打开文件检查是否被占用
        with open(filepath, 'rb') as f:
            f.read(1)
        return True, ""
    except PermissionError:
        return False, f"文件被占用或无权限: {filepath}"
    except Exception as e:
        return False, f"无法访问文件: {e}"


# ============================================================================
# 智能行数检测（优化 #16 - 二分查找版本，大幅提升性能）
# ============================================================================

def _is_row_empty_openpyxl(ws, row_idx: int, max_cols: int) -> bool:
    """检查 openpyxl worksheet 的某一行是否为空"""
    for col_idx in range(1, max_cols + 1):
        val = ws.cell(row=row_idx, column=col_idx).value
        if val is not None and (not isinstance(val, str) or val.strip()):
            return False
    return True


def _detect_actual_data_rows(ws, start_row: int = 1, max_cols: int = None) -> int:
    """
    智能检测实际数据行数，使用二分查找 + 从后向前扫描
    
    优化策略：
    1. 对于小文件（<1000行）：直接使用 max_row
    2. 对于大文件：从后向前扫描，找到最后一个非空行
    
    Args:
        ws: openpyxl worksheet
        start_row: 起始行
        max_cols: 检测的列数，None则使用 max_column
    
    Returns:
        实际数据行数
    """
    max_cols = max_cols or min(ws.max_column or 1, Config.MAX_PREVIEW_COLS)
    max_row = ws.max_row or 1
    
    # 小文件直接返回 max_row（信任 openpyxl 的值）
    if max_row <= 1000:
        return max_row
    
    # 大文件：从后向前扫描，最多检查 200 行
    scan_limit = min(200, max_row - start_row)
    
    for i in range(scan_limit):
        row_idx = max_row - i
        if row_idx < start_row:
            break
        if not _is_row_empty_openpyxl(ws, row_idx, max_cols):
            return row_idx
    
    # 如果后 200 行都是空的，返回 max_row - 200（保守估计）
    return max(start_row, max_row - scan_limit)


def _detect_actual_data_rows_xlrd(ws, start_row: int = 0) -> int:
    """xlrd 版本的实际行数检测（二分查找优化）"""
    max_row = ws.nrows
    max_cols = min(ws.ncols, Config.MAX_PREVIEW_COLS)
    
    # 小文件直接返回
    if max_row <= 1000:
        return max_row
    
    # 从后向前扫描
    scan_limit = min(200, max_row - start_row)
    
    for i in range(scan_limit):
        row_idx = max_row - 1 - i  # xlrd 是 0-based
        if row_idx < start_row:
            break
        row_empty = True
        for col_idx in range(max_cols):
            val = ws.cell_value(row_idx, col_idx)
            if val is not None and val != '':
                row_empty = False
                break
        if not row_empty:
            return row_idx + 1  # 转为 1-based
    
    return max(start_row + 1, max_row - scan_limit)


# ============================================================================
# 合并单元格处理（优化 #5）
# ============================================================================

def _get_merged_cell_value(ws, row: int, col: int, merged_ranges: list = None) -> Any:
    """
    获取单元格值，处理合并单元格的值传播
    
    Args:
        ws: worksheet
        row: 行号 (1-based)
        col: 列号 (1-based)
        merged_ranges: 合并单元格范围列表
    
    Returns:
        单元格值
    """
    if merged_ranges is None:
        try:
            merged_ranges = list(ws.merged_cells.ranges)
        except:
            merged_ranges = []
    
    # 检查是否在合并区域内
    for merged_range in merged_ranges:
        if (merged_range.min_row <= row <= merged_range.max_row and
            merged_range.min_col <= col <= merged_range.max_col):
            # 返回合并区域左上角的值
            return ws.cell(row=merged_range.min_row, column=merged_range.min_col).value
    
    return ws.cell(row=row, column=col).value


# ============================================================================
# 文件信息获取（快速，不读数据）
# ============================================================================

def get_file_info(filepath: str) -> Dict[str, Any]:
    """
    快速获取文件基本信息（不读取数据内容）
    
    Returns:
        {
            'success': True/False,
            'data': {
                'name': 文件名,
                'type': 文件类型(excel/word/pdf/csv),
                'size_bytes': 文件大小,
                'needs_convert': 是否需要转换格式,
                'structure': {  # 结构信息
                    'sheets': [{'name': 表名, 'rows': 行数, 'cols': 列数, 'actual_rows': 实际行数}],
                    'paragraphs_count': 段落数,  # Word
                    'pages_count': 页数,  # PDF
                }
            }
        }
    """
    filepath = Path(filepath)
    
    # 检查文件可访问性
    accessible, error_msg = _check_file_accessible(filepath)
    if not accessible:
        return _make_response(False, error=error_msg)
    
    file_type, needs_convert = get_file_type(str(filepath))
    
    info = {
        'name': filepath.name,
        'path': str(filepath.absolute()),
        'type': file_type,
        'size_bytes': filepath.stat().st_size,
        'size_human': _human_size(filepath.stat().st_size),
        'needs_convert': needs_convert,
        'structure': {}
    }
    
    try:
        if file_type == 'excel':
            info['structure'] = _get_excel_structure(filepath, needs_convert)
        elif file_type == 'word':
            info['structure'] = _get_word_structure(filepath, needs_convert)
        elif file_type == 'pdf':
            info['structure'] = _get_pdf_structure(filepath)
        elif file_type == 'csv':
            info['structure'] = _get_csv_structure(filepath)
    except Exception as e:
        info['structure_error'] = str(e)
        logger.error(f"获取文件结构失败: {filepath}, 错误: {e}")
    
    return _make_response(True, data=info)


def _get_excel_structure(filepath: Path, needs_convert: bool) -> Dict:
    """获取Excel结构信息（优化 #2: 只读取必要元数据）"""
    if needs_convert:
        import xlrd
        wb = xlrd.open_workbook(str(filepath), on_demand=True)  # 按需加载
        sheets = []
        for name in wb.sheet_names():
            ws = wb.sheet_by_name(name)
            actual_rows = _detect_actual_data_rows_xlrd(ws)
            sheets.append({
                'name': name, 
                'rows': ws.nrows, 
                'cols': ws.ncols,
                'actual_rows': actual_rows
            })
            wb.unload_sheet(name)  # 释放内存
        return {'sheets': sheets, 'sheet_count': len(sheets)}
    else:
        from openpyxl import load_workbook
        # 只读模式，不加载数据
        wb = load_workbook(filepath, read_only=True, data_only=True)
        sheets = []
        for name in wb.sheetnames:
            ws = wb[name]
            actual_rows = _detect_actual_data_rows(ws)
            sheets.append({
                'name': name, 
                'rows': ws.max_row, 
                'cols': ws.max_column,
                'actual_rows': actual_rows
            })
        wb.close()
        return {'sheets': sheets, 'sheet_count': len(sheets)}


def _get_word_structure(filepath: Path, needs_convert: bool) -> Dict:
    """获取Word结构信息"""
    if needs_convert:
        return {'note': '旧格式Word，需转换后读取', 'needs_convert': True}
    from docx import Document
    doc = Document(filepath)
    return {
        'paragraphs_count': len(doc.paragraphs),
        'tables_count': len(doc.tables),
        'has_tables': len(doc.tables) > 0
    }


def _get_pdf_structure(filepath: Path) -> Dict:
    """获取PDF结构信息"""
    import pdfplumber
    with pdfplumber.open(filepath) as pdf:
        return {'pages_count': len(pdf.pages)}


def _get_csv_structure(filepath: Path) -> Dict:
    """获取CSV结构信息（新增 #19）"""
    import chardet
    
    # 检测编码
    with open(filepath, 'rb') as f:
        raw = f.read(10000)
        detected = chardet.detect(raw)
        encoding = detected.get('encoding', 'utf-8')
    
    # 读取行数和列数
    with open(filepath, 'r', encoding=encoding, errors='ignore') as f:
        reader = csv.reader(f)
        first_row = next(reader, [])
        row_count = sum(1 for _ in reader) + 1
    
    return {
        'rows': row_count,
        'cols': len(first_row),
        'encoding': encoding,
        'headers': first_row[:Config.MAX_PREVIEW_COLS]
    }


# ============================================================================
# Excel 读取（优化 #1, #5, #11）
# ============================================================================

def _read_excel_common(
    filepath: Path,
    sheet_name: str = None,
    preview_rows: int = None,
    header_rows: int = 1,
    handle_merged: bool = True,
    progress_callback: Callable[[int, int], None] = None
) -> Dict[str, Any]:
    """
    Excel 读取的公共逻辑（优化 #11: 抽取公共逻辑）
    
    Args:
        filepath: 文件路径
        sheet_name: 工作表名，None则读取所有
        preview_rows: 只读取前N行数据（None=智能检测实际行数）
        header_rows: 表头占用的行数
        handle_merged: 是否处理合并单元格
        progress_callback: 进度回调函数 (current, total)
    
    Returns:
        标准化的数据结构
    """
    file_type, needs_convert = get_file_type(str(filepath))
    
    if needs_convert:
        return _read_excel_xlrd_impl(filepath, sheet_name, preview_rows, header_rows, progress_callback)
    else:
        return _read_excel_openpyxl_impl(filepath, sheet_name, preview_rows, header_rows, handle_merged, progress_callback)


def _read_excel_openpyxl_impl(
    filepath: Path,
    sheet_name: str = None,
    preview_rows: int = None,
    header_rows: int = 1,
    handle_merged: bool = True,
    progress_callback: Callable = None
) -> Dict[str, Any]:
    """openpyxl 实现"""
    from openpyxl import load_workbook
    
    result = {'file': filepath.name, 'sheets': []}
    
    try:
        wb = load_workbook(filepath, read_only=True, data_only=True)
        sheets_to_read = [sheet_name] if sheet_name else wb.sheetnames
        
        for sheet_idx, sn in enumerate(sheets_to_read):
            if sn not in wb.sheetnames:
                continue
            ws = wb[sn]
            
            # 智能检测实际行数（优化 #1, #16）
            actual_rows = _detect_actual_data_rows(ws, header_rows + 1)
            data_rows_count = max(0, actual_rows - header_rows)
            
            # 确定读取范围
            if preview_rows is not None:
                rows_to_read = min(preview_rows, data_rows_count)
            else:
                rows_to_read = data_rows_count
            
            sheet_data = {
                'name': sn,
                'rows': ws.max_row,
                'actual_rows': actual_rows,
                'cols': ws.max_column,
                'header_rows': header_rows,
                'headers': [],
                'headers_raw': [],
                'data_rows': data_rows_count,
                'data': [],
                'preview_only': preview_rows is not None,
                'column_types': {}
            }
            
            # 获取合并单元格信息
            merged_ranges = []
            if handle_merged:
                try:
                    merged_ranges = list(ws.merged_cells.ranges)
                except:
                    pass
            
            max_cols = min(ws.max_column or 1, Config.MAX_PREVIEW_COLS)
            
            # 读取原始表头（处理合并单元格）
            for row_idx in range(1, header_rows + 1):
                row_headers = []
                for col_idx in range(1, max_cols + 1):
                    if handle_merged and merged_ranges:
                        val = _get_merged_cell_value(ws, row_idx, col_idx, merged_ranges)
                    else:
                        val = ws.cell(row=row_idx, column=col_idx).value
                    row_headers.append(str(val).strip() if val else '')
                sheet_data['headers_raw'].append(row_headers)
            
            # 合并表头
            sheet_data['headers'] = _merge_headers(sheet_data['headers_raw'], max_cols)
            
            # 读取数据（优化 #1: 智能停止）
            data_start = header_rows + 1
            data_end = data_start + rows_to_read
            empty_count = 0
            
            for row_idx, row in enumerate(ws.iter_rows(min_row=data_start, max_row=data_end, 
                                                        max_col=max_cols, values_only=True)):
                if progress_callback:
                    progress_callback(row_idx + 1, rows_to_read)
                
                row_list = list(row)
                
                # 检查空行
                if _is_row_empty(row_list):
                    empty_count += 1
                    if empty_count >= Config.EMPTY_ROW_THRESHOLD:
                        break
                else:
                    empty_count = 0
                    sheet_data['data'].append(row_list)
            
            # 推断列类型（优化 #7: 增加采样量）
            sheet_data['column_types'] = _infer_column_types(
                sheet_data['headers'], 
                sheet_data['data'],
                sample_size=Config.TYPE_INFER_SAMPLE_SIZE
            )
            
            result['sheets'].append(sheet_data)
        
        wb.close()
        return _make_response(True, data=result)
        
    except Exception as e:
        logger.error(f"读取Excel失败: {filepath}, 错误: {e}")
        return _make_response(False, error=str(e))


def _read_excel_xlrd_impl(
    filepath: Path,
    sheet_name: str = None,
    preview_rows: int = None,
    header_rows: int = 1,
    progress_callback: Callable = None
) -> Dict[str, Any]:
    """xlrd 实现（旧版 .xls 文件）"""
    import xlrd
    
    result = {'file': Path(filepath).name, 'sheets': []}
    
    try:
        wb = xlrd.open_workbook(str(filepath))
        sheets_to_read = [sheet_name] if sheet_name else wb.sheet_names()
        
        for sn in sheets_to_read:
            ws = wb.sheet_by_name(sn)
            
            # 智能检测实际行数
            actual_rows = _detect_actual_data_rows_xlrd(ws, header_rows)
            data_rows_count = max(0, actual_rows - header_rows)
            
            if preview_rows is not None:
                rows_to_read = min(preview_rows, data_rows_count)
            else:
                rows_to_read = data_rows_count
            
            max_cols = min(ws.ncols, Config.MAX_PREVIEW_COLS)
            
            sheet_data = {
                'name': sn,
                'rows': ws.nrows,
                'actual_rows': actual_rows,
                'cols': ws.ncols,
                'header_rows': header_rows,
                'headers': [],
                'headers_raw': [],
                'data_rows': data_rows_count,
                'data': [],
                'preview_only': preview_rows is not None,
                'column_types': {}
            }
            
            # 读取表头
            for row_idx in range(header_rows):
                row_headers = [
                    str(ws.cell_value(row_idx, c)).strip() if ws.cell_value(row_idx, c) else '' 
                    for c in range(max_cols)
                ]
                sheet_data['headers_raw'].append(row_headers)
            
            # 合并表头
            sheet_data['headers'] = _merge_headers(sheet_data['headers_raw'], max_cols)
            
            # 读取数据
            empty_count = 0
            for row_idx in range(header_rows, header_rows + rows_to_read):
                if row_idx >= ws.nrows:
                    break
                    
                if progress_callback:
                    progress_callback(row_idx - header_rows + 1, rows_to_read)
                
                row_data = [ws.cell_value(row_idx, c) for c in range(max_cols)]
                
                if _is_row_empty(row_data):
                    empty_count += 1
                    if empty_count >= Config.EMPTY_ROW_THRESHOLD:
                        break
                else:
                    empty_count = 0
                    sheet_data['data'].append(row_data)
            
            sheet_data['column_types'] = _infer_column_types(
                sheet_data['headers'], 
                sheet_data['data'],
                sample_size=Config.TYPE_INFER_SAMPLE_SIZE
            )
            result['sheets'].append(sheet_data)
            
        return _make_response(True, data=result)
        
    except Exception as e:
        logger.error(f"读取Excel(xls)失败: {filepath}, 错误: {e}")
        return _make_response(False, error=str(e))


def _merge_headers(headers_raw: List[List[str]], max_cols: int) -> List[str]:
    """合并多行表头"""
    if not headers_raw:
        return [f'列{i+1}' for i in range(max_cols)]
    
    merged = []
    for col_idx in range(max_cols):
        header_parts = []
        for row in headers_raw:
            if col_idx < len(row) and row[col_idx]:
                header_parts.append(row[col_idx])
        merged.append(' > '.join(header_parts) if header_parts else f'列{col_idx+1}')
    
    return merged


def read_excel(filepath: str, sheet_name: str = None, 
               preview_rows: int = None, header_rows: int = 1,
               handle_merged: bool = True,
               progress_callback: Callable[[int, int], None] = None) -> Dict[str, Any]:
    """
    读取Excel文件，返回标准化结构
    
    Args:
        filepath: 文件路径
        sheet_name: 工作表名，None则读取所有
        preview_rows: 只读取前N行数据（None=智能检测实际行数）
        header_rows: 表头占用的行数
        handle_merged: 是否处理合并单元格值传播
        progress_callback: 进度回调函数 (current, total)
    
    Returns:
        标准化的Excel数据结构
    """
    filepath = Path(filepath)
    
    # 检查文件可访问性
    accessible, error_msg = _check_file_accessible(filepath)
    if not accessible:
        return _make_response(False, error=error_msg)
    
    return _read_excel_common(filepath, sheet_name, preview_rows, header_rows, 
                              handle_merged, progress_callback)


# ============================================================================
# 类型推断（优化 #4, #7）
# ============================================================================

def _infer_column_types(headers: List[str], data: List[List], 
                        sample_size: int = 50) -> Dict[str, str]:
    """
    推断列的数据类型（优化 #4: 使用预编译正则，#7: 增加采样量）
    
    Args:
        headers: 表头列表
        data: 数据行列表
        sample_size: 采样行数
    
    Returns:
        {列名: 类型} 字典
    """
    types = {}
    if not data:
        return types
    
    # 智能采样：均匀分布采样
    total_rows = len(data)
    if total_rows <= sample_size:
        sample_indices = range(total_rows)
    else:
        step = total_rows / sample_size
        sample_indices = [int(i * step) for i in range(sample_size)]
    
    for i, header in enumerate(headers):
        if i >= len(data[0]) if data else 0:
            continue
        
        # 采样
        samples = []
        for idx in sample_indices:
            if idx < len(data) and i < len(data[idx]):
                val = data[idx][i]
                if val is not None:
                    samples.append(val)
        
        if not samples:
            types[header] = 'empty'
            continue
        
        # 类型统计
        type_counts = {'numeric': 0, 'date': 0, 'percentage': 0, 'text': 0}
        
        for v in samples:
            if isinstance(v, (int, float)):
                type_counts['numeric'] += 1
            elif isinstance(v, datetime):
                type_counts['date'] += 1
            elif isinstance(v, str):
                v_stripped = v.strip()
                if RE_NUMERIC.match(v_stripped):
                    type_counts['numeric'] += 1
                elif RE_DATE.match(v_stripped):
                    type_counts['date'] += 1
                elif RE_PERCENTAGE.match(v_stripped):
                    type_counts['percentage'] += 1
                else:
                    type_counts['text'] += 1
            else:
                type_counts['text'] += 1
        
        # 确定主要类型（超过50%）
        total = len(samples)
        if type_counts['date'] > total * 0.5:
            types[header] = 'datetime'
        elif type_counts['percentage'] > total * 0.3:
            types[header] = 'percentage'
        elif type_counts['numeric'] > total * 0.5:
            types[header] = 'numeric'
        else:
            types[header] = 'text'
    
    return types


# ============================================================================
# Excel 写入（优化 #9, #20）
# ============================================================================

def write_excel(data: List[List], filepath: str, headers: List[str] = None, 
                sheet_name: str = 'Sheet1', 
                template_path: str = None,
                use_write_only: bool = False) -> Dict[str, Any]:
    """
    写入Excel文件
    
    Args:
        data: 二维数据列表
        filepath: 输出文件路径
        headers: 表头列表
        sheet_name: 工作表名称
        template_path: 模板文件路径（优化 #9: 支持模板写入）
        use_write_only: 是否使用只写模式（优化 #20: 大数据量优化）
    
    Returns:
        {'success': True/False, 'data': {'path': 保存路径, 'rows': 行数}}
    """
    from openpyxl import Workbook, load_workbook
    
    try:
        if template_path and Path(template_path).exists():
            # 基于模板写入
            wb = load_workbook(template_path)
            if sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
            else:
                ws = wb.active
        elif use_write_only and len(data) > 1000:
            # 大数据量使用只写模式
            wb = Workbook(write_only=True)
            ws = wb.create_sheet(sheet_name)
        else:
            wb = Workbook()
            ws = wb.active
            ws.title = sheet_name
        
        if headers:
            ws.append(headers)
        
        for row in data:
            ws.append(row)
        
        wb.save(filepath)
        
        return _make_response(True, data={
            'path': str(Path(filepath).absolute()),
            'rows': len(data) + (1 if headers else 0)
        })
    except Exception as e:
        logger.error(f"写入Excel失败: {filepath}, 错误: {e}")
        return _make_response(False, error=str(e))


# ============================================================================
# CSV 处理（新增 #19）
# ============================================================================

def read_csv(filepath: str, preview_rows: int = None, 
             encoding: str = None, delimiter: str = None) -> Dict[str, Any]:
    """
    读取CSV文件
    
    Args:
        filepath: 文件路径
        preview_rows: 预览行数
        encoding: 编码，None则自动检测
        delimiter: 分隔符，None则自动检测
    
    Returns:
        标准化的数据结构
    """
    import chardet
    
    filepath = Path(filepath)
    
    accessible, error_msg = _check_file_accessible(filepath)
    if not accessible:
        return _make_response(False, error=error_msg)
    
    try:
        # 自动检测编码
        if encoding is None:
            with open(filepath, 'rb') as f:
                raw = f.read(10000)
                detected = chardet.detect(raw)
                encoding = detected.get('encoding', 'utf-8')
        
        # 自动检测分隔符
        if delimiter is None:
            with open(filepath, 'r', encoding=encoding, errors='ignore') as f:
                sample = f.read(2000)
                sniffer = csv.Sniffer()
                try:
                    dialect = sniffer.sniff(sample)
                    delimiter = dialect.delimiter
                except:
                    delimiter = ','
        
        # 读取数据
        with open(filepath, 'r', encoding=encoding, errors='ignore') as f:
            reader = csv.reader(f, delimiter=delimiter)
            headers = next(reader, [])
            
            data = []
            for i, row in enumerate(reader):
                if preview_rows and i >= preview_rows:
                    break
                data.append(row)
        
        result = {
            'file': filepath.name,
            'encoding': encoding,
            'delimiter': delimiter,
            'headers': headers,
            'data': data,
            'rows': len(data),
            'cols': len(headers),
            'preview_only': preview_rows is not None
        }
        
        return _make_response(True, data=result)
        
    except Exception as e:
        logger.error(f"读取CSV失败: {filepath}, 错误: {e}")
        return _make_response(False, error=str(e))


def write_csv(data: List[List], filepath: str, headers: List[str] = None,
              encoding: str = 'utf-8-sig', delimiter: str = ',') -> Dict[str, Any]:
    """
    写入CSV文件
    
    Args:
        data: 二维数据列表
        filepath: 输出文件路径
        headers: 表头列表
        encoding: 编码（默认utf-8-sig，Excel友好）
        delimiter: 分隔符
    
    Returns:
        标准化响应
    """
    try:
        with open(filepath, 'w', encoding=encoding, newline='') as f:
            writer = csv.writer(f, delimiter=delimiter)
            if headers:
                writer.writerow(headers)
            writer.writerows(data)
        
        return _make_response(True, data={
            'path': str(Path(filepath).absolute()),
            'rows': len(data) + (1 if headers else 0)
        })
    except Exception as e:
        logger.error(f"写入CSV失败: {filepath}, 错误: {e}")
        return _make_response(False, error=str(e))


# ============================================================================
# Word 处理（优化 #8）
# ============================================================================

def read_word(filepath: str, include_formatting: bool = False) -> Dict[str, Any]:
    """
    读取Word文档，返回标准化结构
    
    Args:
        filepath: 文件路径
        include_formatting: 是否包含格式信息
    
    Returns:
        标准化的Word文档结构
    """
    from docx import Document
    
    filepath = Path(filepath)
    
    accessible, error_msg = _check_file_accessible(filepath)
    if not accessible:
        return _make_response(False, error=error_msg)
    
    file_type, needs_convert = get_file_type(str(filepath))
    
    if needs_convert:
        new_path = convert_to_new_format(str(filepath))
        if new_path:
            filepath = new_path
        else:
            # 优化 #8: 更友好的错误提示
            return _make_response(False, error=(
                f'无法自动转换旧格式Word文件 ({filepath.suffix})。\n'
                f'请尝试以下方法：\n'
                f'1. 安装 LibreOffice: brew install --cask libreoffice\n'
                f'2. 或手动用WPS/Word打开文件，另存为 .docx 格式'
            ))
    
    try:
        doc = Document(filepath)
        
        result = {
            'file': filepath.name,
            'structure': {
                'paragraphs_count': len(doc.paragraphs),
                'tables_count': len(doc.tables),
                'headings_count': 0
            },
            'content': {
                'paragraphs': [],
                'headings': [],
                'tables': []
            },
            'outline': []
        }
        
        # 解析段落
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            para_info = {'text': text, 'style': para.style.name}
            
            # 识别标题
            if para.style.name.startswith('Heading'):
                try:
                    level = int(para.style.name.replace('Heading ', '').replace('Heading', '1'))
                except:
                    level = 1
                para_info['level'] = level
                result['content']['headings'].append({'text': text, 'level': level})
                result['outline'].append(('  ' * (level - 1)) + text)
                result['structure']['headings_count'] += 1
            
            result['content']['paragraphs'].append(para_info)
        
        # 解析表格
        for table in doc.tables:
            table_info = {
                'rows': len(table.rows),
                'cols': len(table.columns),
                'headers': [],
                'data': []
            }
            
            for i, row in enumerate(table.rows):
                row_data = [cell.text.strip() for cell in row.cells]
                if i == 0:
                    table_info['headers'] = row_data
                else:
                    table_info['data'].append(row_data)
            
            result['content']['tables'].append(table_info)
        
        return _make_response(True, data=result)
        
    except Exception as e:
        logger.error(f"读取Word失败: {filepath}, 错误: {e}")
        return _make_response(False, error=str(e))


def write_word(content: Dict, filepath: str) -> Dict[str, Any]:
    """
    写入Word文档
    
    Args:
        content: {
            'title': 标题,
            'paragraphs': [文本] 或 [{'text': 文本, 'style': 样式}],
            'tables': [{'headers': [], 'data': [[]]}]
        }
    
    Returns:
        标准化响应
    """
    from docx import Document
    
    try:
        doc = Document()
        
        if 'title' in content:
            doc.add_heading(content['title'], level=1)
        
        for para in content.get('paragraphs', []):
            if isinstance(para, str):
                doc.add_paragraph(para)
            elif isinstance(para, dict):
                p = doc.add_paragraph(para.get('text', ''))
        
        for table_info in content.get('tables', []):
            headers = table_info.get('headers', [])
            data = table_info.get('data', [])
            
            if not headers and not data:
                continue
            
            rows = len(data) + (1 if headers else 0)
            cols = len(headers) if headers else (len(data[0]) if data else 0)
            
            if cols > 0:
                table = doc.add_table(rows=rows, cols=cols)
                table.style = 'Table Grid'
                
                if headers:
                    for i, h in enumerate(headers):
                        table.cell(0, i).text = str(h)
                
                start_row = 1 if headers else 0
                for r_idx, row in enumerate(data):
                    for c_idx, val in enumerate(row[:cols]):
                        table.cell(start_row + r_idx, c_idx).text = str(val) if val else ''
        
        doc.save(filepath)
        return _make_response(True, data={'path': str(Path(filepath).absolute())})
        
    except Exception as e:
        logger.error(f"写入Word失败: {filepath}, 错误: {e}")
        return _make_response(False, error=str(e))


# ============================================================================
# PDF 处理
# ============================================================================

def read_pdf(filepath: str, pages: List[int] = None, 
             extract_tables: bool = True) -> Dict[str, Any]:
    """
    读取PDF文件，返回标准化结构
    
    Args:
        filepath: 文件路径
        pages: 要读取的页码列表（从0开始），None=全部
        extract_tables: 是否提取表格
    
    Returns:
        标准化的PDF内容结构
    """
    import pdfplumber
    
    filepath = Path(filepath)
    
    accessible, error_msg = _check_file_accessible(filepath)
    if not accessible:
        return _make_response(False, error=error_msg)
    
    try:
        result = {
            'file': filepath.name,
            'structure': {'pages_count': 0, 'has_tables': False},
            'pages': []
        }
        
        with pdfplumber.open(filepath) as pdf:
            result['structure']['pages_count'] = len(pdf.pages)
            pages_to_read = pages if pages else range(len(pdf.pages))
            
            for page_num in pages_to_read:
                if page_num >= len(pdf.pages):
                    continue
                    
                page = pdf.pages[page_num]
                page_data = {
                    'page_num': page_num + 1,
                    'text': page.extract_text() or '',
                    'tables': []
                }
                
                if extract_tables:
                    tables = page.extract_tables() or []
                    for table in tables:
                        if table:
                            page_data['tables'].append({
                                'rows': len(table),
                                'cols': len(table[0]) if table else 0,
                                'data': table
                            })
                            result['structure']['has_tables'] = True
                
                result['pages'].append(page_data)
        
        return _make_response(True, data=result)
        
    except Exception as e:
        logger.error(f"读取PDF失败: {filepath}, 错误: {e}")
        return _make_response(False, error=str(e))


# ============================================================================
# Excel 深度分析（优化 #6, #13, #15）
# ============================================================================

# 配置常量（优化 #15）
ROLE_PATTERNS = {
    '汇总表': ['汇总', '统计', '合计', '总计', '概览', '总表'],
    '明细表': ['明细', '详情', '清单', '列表', '记录', '数据'],
    '模板': ['模板', '样表', '表样', '格式', '填报'],
    '配置表': ['配置', '参数', '字典', '编码', '对照'],
    '线索表': ['线索', '疑点', '问题', '异常', '风险'],
}

FIELD_TYPES = {
    '标识字段': ['编号', '编码', 'ID', 'id', '代码', '流水号', '序号'],
    '名称字段': ['名称', '姓名', '机构', '单位', '药品', '药店'],
    '时间字段': ['时间', '日期', '年', '月', '日'],
    '金额字段': ['金额', '费用', '价格', '元', '万元', '合计'],
    '数量字段': ['数量', '次数', '人数', '家数', '个数'],
    '状态字段': ['状态', '结果', '情况', '等级', '类型', '是否'],
    '地区字段': ['省', '市', '区', '县', '地区', '地址'],
}


def _detect_header_rows(rows_data: List[List[str]]) -> int:
    """
    智能检测表头行数（优化 #6）
    
    策略：
    1. 检查每行的数值比例
    2. 检查是否有合并单元格特征（连续空值）
    3. 检查行内容的重复性
    """
    if not rows_data or len(rows_data) < 2:
        return 1
    
    for i, row in enumerate(rows_data):
        if i == 0:
            continue
        
        non_empty = [v for v in row if v and str(v).strip()]
        if not non_empty:
            continue
        
        # 计算数值比例
        numeric_count = sum(1 for v in non_empty if RE_NUMERIC.match(str(v).strip()))
        numeric_ratio = numeric_count / len(non_empty) if non_empty else 0
        
        # 如果数值比例超过50%，认为是数据行
        if numeric_ratio > 0.5:
            return i
        
        # 检查是否包含典型的数据特征（日期、金额等）
        has_date = any(RE_DATE.match(str(v).strip()) for v in non_empty if v)
        has_currency = any(RE_CURRENCY.match(str(v).strip()) for v in non_empty if v)
        
        if has_date or has_currency:
            return i
    
    return min(len(rows_data), Config.MAX_HEADER_ROWS)


def _analyze_sheet_role(sheet_name: str, data_rows: int, headers: List[str]) -> str:
    """分析工作表角色"""
    # 根据名称判断
    for role, keywords in ROLE_PATTERNS.items():
        if any(kw in sheet_name for kw in keywords):
            return role
    
    # 根据数据量判断
    if data_rows <= 2:
        return '模板/汇总表'
    elif data_rows > 100:
        return '明细表'
    else:
        return '数据表'


def _analyze_field(header: str) -> Dict[str, Any]:
    """分析单个字段"""
    field_info = {'name': header, 'type': '其他', 'importance': 'normal'}
    
    for field_type, keywords in FIELD_TYPES.items():
        if any(kw in header for kw in keywords):
            field_info['type'] = field_type
            if field_type in ['标识字段', '金额字段', '时间字段']:
                field_info['importance'] = 'high'
            break
    
    return field_info


def _analyze_sheet(ws, sheet_name: str, wb_merged_ranges: dict = None) -> Dict[str, Any]:
    """
    分析单个工作表（优化 #13: 拆分子函数）
    """
    sheet_info = {
        'name': sheet_name,
        'rows': ws.max_row,
        'cols': ws.max_column,
        'role': '未知',
        'header_rows': 1,
        'headers': [],
        'header_structure': [],
        'data_preview': [],
        'field_analysis': [],
        'data_stats': {},
    }
    
    max_cols = min(ws.max_column or 1, Config.MAX_PREVIEW_COLS)
    
    # 读取前几行用于分析
    header_rows_data = []
    for row_idx in range(1, min(Config.MAX_HEADER_ROWS + 1, (ws.max_row or 0) + 1)):
        row_data = []
        for col_idx in range(1, max_cols + 1):
            val = ws.cell(row=row_idx, column=col_idx).value
            row_data.append(str(val).strip() if val else '')
        header_rows_data.append(row_data)
    
    # 智能检测表头行数
    header_row_count = _detect_header_rows(header_rows_data)
    sheet_info['header_rows'] = header_row_count
    sheet_info['header_structure'] = header_rows_data[:header_row_count]
    
    # 合并表头
    sheet_info['headers'] = _merge_headers(header_rows_data[:header_row_count], max_cols)
    
    # 检测实际数据行数
    actual_rows = _detect_actual_data_rows(ws, header_row_count + 1, max_cols)
    data_rows = max(0, actual_rows - header_row_count)
    
    # 数据预览
    data_start = header_row_count + 1
    preview_count = 0
    for row_idx in range(data_start, min(data_start + 5, actual_rows + 1)):
        row_data = []
        for col_idx in range(1, min(max_cols + 1, 15)):
            val = ws.cell(row=row_idx, column=col_idx).value
            if val is not None:
                val_str = str(val)[:50]
                row_data.append(val_str)
            else:
                row_data.append('')
        if any(row_data):
            sheet_info['data_preview'].append(row_data)
            preview_count += 1
    
    # 字段分析
    for header in sheet_info['headers'][:20]:
        sheet_info['field_analysis'].append(_analyze_field(header))
    
    # 角色判断
    sheet_info['role'] = _analyze_sheet_role(sheet_name, data_rows, sheet_info['headers'])
    
    # 数据统计
    sheet_info['data_stats'] = {
        'data_rows': data_rows,
        'actual_rows': actual_rows,
        'is_empty': data_rows <= 0,
        'is_template': data_rows <= 2 and ('模板' in sheet_name or '表样' in sheet_name),
    }
    
    return sheet_info


def analyze_excel_structure(filepath: str, 
                           progress_callback: Callable[[str, int, int], None] = None) -> Dict[str, Any]:
    """
    深度分析Excel结构，用于理解任务意图
    
    Args:
        filepath: Excel文件路径
        progress_callback: 进度回调 (sheet_name, current, total)
    
    Returns:
        详细的结构分析结果
    """
    from openpyxl import load_workbook
    
    filepath = Path(filepath)
    
    accessible, error_msg = _check_file_accessible(filepath)
    if not accessible:
        return {'error': error_msg}
    
    result = {
        'file': filepath.name,
        'overview': {},
        'sheets': [],
        'task_hints': [],
        'key_fields': [],
    }
    
    try:
        wb = load_workbook(filepath, read_only=True, data_only=True)
        total_sheets = len(wb.sheetnames)
        total_rows = 0
        total_data_rows = 0
        
        for idx, sheet_name in enumerate(wb.sheetnames):
            if progress_callback:
                progress_callback(sheet_name, idx + 1, total_sheets)
            
            ws = wb[sheet_name]
            sheet_info = _analyze_sheet(ws, sheet_name)
            
            total_rows += sheet_info['rows']
            total_data_rows += sheet_info['data_stats'].get('data_rows', 0)
            
            # 收集关键字段
            for field in sheet_info['field_analysis']:
                if field['importance'] == 'high':
                    result['key_fields'].append(field['name'])
            
            result['sheets'].append(sheet_info)
        
        wb.close()
        
        # 文件概览
        result['overview'] = {
            'sheet_count': total_sheets,
            'total_rows': total_rows,
            'total_data_rows': total_data_rows,
            'has_template': any(s['data_stats'].get('is_template') for s in result['sheets']),
            'has_detail': any(s['data_stats'].get('data_rows', 0) > 50 for s in result['sheets']),
        }
        
        # 任务意图推断
        result['task_hints'] = _infer_task_hints(result)
            
    except Exception as e:
        result['error'] = str(e)
        logger.error(f"分析Excel结构失败: {filepath}, 错误: {e}")
    
    return result


def _infer_task_hints(result: Dict) -> List[str]:
    """推断任务意图"""
    hints = []
    
    if result['overview'].get('has_template') and result['overview'].get('has_detail'):
        hints.append('可能需要：根据明细数据填写模板/汇总表')
    
    if any('线索' in s.get('role', '') for s in result.get('sheets', [])):
        hints.append('可能需要：核查线索、分析问题')
    
    if result['overview'].get('total_data_rows', 0) > 500:
        hints.append('数据量较大，可能需要：数据筛选、统计分析')
    
    if result.get('key_fields'):
        unique_fields = list(dict.fromkeys(result['key_fields'][:10]))
        hints.append(f'关键字段：{", ".join(unique_fields)}')
    
    return hints


def print_excel_analysis(result: Dict):
    """打印Excel分析结果"""
    print(f"\n{'='*70}")
    print(f"📊 Excel深度分析: {result.get('file', 'Unknown')}")
    print('='*70)
    
    if 'error' in result:
        print(f"❌ 错误: {result['error']}")
        return
    
    ov = result.get('overview', {})
    print(f"\n📋 文件概览")
    print(f"   Sheet数量: {ov.get('sheet_count', 0)}")
    print(f"   总数据行: {ov.get('total_data_rows', 0)}")
    
    if result.get('task_hints'):
        print(f"\n💡 任务意图推断")
        for hint in result['task_hints']:
            print(f"   • {hint}")
    
    for sheet in result.get('sheets', []):
        print(f"\n{'─'*70}")
        print(f"📑 Sheet: {sheet['name']}")
        print(f"   角色: {sheet['role']}")
        print(f"   大小: {sheet['rows']}行 × {sheet['cols']}列 (实际数据: {sheet['data_stats'].get('data_rows', 0)}行)")
        print(f"   表头行数: {sheet['header_rows']}")
        
        if sheet.get('headers'):
            print(f"\n   字段列表:")
            for i, h in enumerate(sheet['headers'][:15], 1):
                field_info = sheet['field_analysis'][i-1] if i <= len(sheet['field_analysis']) else {}
                importance = '⭐' if field_info.get('importance') == 'high' else ''
                field_type = f"[{field_info.get('type', '')}]" if field_info.get('type') != '其他' else ''
                print(f"   {i:2}. {h} {importance} {field_type}")
        
        if sheet.get('data_preview'):
            print(f"\n   数据预览 (前{len(sheet['data_preview'])}行):")
            for row in sheet['data_preview'][:3]:
                preview = ' | '.join(str(v)[:15] for v in row[:6])
                print(f"      {preview}...")


# ============================================================================
# 目录分析（优化 #3: 可选并行处理）
# ============================================================================

def analyze_directory(directory: str = '.', preview_rows: int = 5, 
                      verbose: bool = True,
                      parallel: bool = False) -> Dict[str, Any]:
    """
    分析目录下所有Office文件，返回标准化结构供AI理解
    
    Args:
        directory: 目录路径
        preview_rows: 每个文件预览的数据行数
        verbose: 是否打印详情
        parallel: 是否并行处理（多文件时更快）
    
    Returns:
        标准化的目录分析结果
    """
    directory = Path(directory)
    
    result = {
        'directory': str(directory.absolute()),
        'analyzed_at': datetime.now().isoformat(),
        'summary': {
            'total_files': 0,
            'by_type': {},
            'total_data_rows': 0,
            'keywords': [],
            'work_type_guess': {}
        },
        'files': []
    }
    
    files = scan_office_files(str(directory))
    result['summary']['total_files'] = len(files)
    
    if verbose:
        print(f"发现 {len(files)} 个Office文件")
    
    all_keywords = []
    
    # 处理函数
    def process_file(f: Path) -> Tuple[Dict, List[str]]:
        file_keywords = []
        file_type, needs_convert = get_file_type(str(f))
        
        file_info = {
            'name': f.name,
            'type': file_type,
            'size': _human_size(f.stat().st_size),
            'needs_convert': needs_convert,
            'structure': {},
            'preview': {}
        }
        
        data_rows = 0
        
        if file_type == 'excel':
            read_result = read_excel(str(f), preview_rows=preview_rows)
            if read_result['success']:
                data = read_result['data']
                file_info['structure'] = {
                    'sheets': [{'name': s['name'], 'rows': s['rows'], 
                               'actual_rows': s.get('actual_rows', s['rows']),
                               'cols': s['cols'], 'headers': s['headers']} 
                              for s in data['sheets']]
                }
                file_info['preview'] = {
                    'sheets': [{'name': s['name'], 'data': s['data'][:preview_rows]} 
                              for s in data['sheets']]
                }
                for s in data['sheets']:
                    data_rows += s.get('data_rows', 0)
                    file_keywords.extend(s['headers'])
        
        elif file_type == 'word':
            read_result = read_word(str(f))
            if read_result['success']:
                data = read_result['data']
                file_info['structure'] = data['structure']
                file_info['preview'] = {
                    'outline': data['outline'][:20],
                    'paragraphs': [p['text'][:100] for p in data['content']['paragraphs'][:10]]
                }
                file_keywords.extend([h['text'] for h in data['content']['headings']])
        
        elif file_type == 'pdf':
            read_result = read_pdf(str(f), pages=[0])
            if read_result['success']:
                data = read_result['data']
                file_info['structure'] = data['structure']
                if data['pages']:
                    file_info['preview'] = {'first_page_text': data['pages'][0]['text'][:500]}
        
        elif file_type == 'csv':
            read_result = read_csv(str(f), preview_rows=preview_rows)
            if read_result['success']:
                data = read_result['data']
                file_info['structure'] = {
                    'rows': data['rows'],
                    'cols': data['cols'],
                    'encoding': data['encoding']
                }
                file_info['preview'] = {'headers': data['headers'], 'data': data['data'][:5]}
                file_keywords.extend(data['headers'])
                data_rows = data['rows']
        
        return file_info, file_keywords, data_rows, file_type
    
    # 处理文件
    if parallel and len(files) > 3:
        # 并行处理（优化 #3）
        from concurrent.futures import ThreadPoolExecutor, as_completed
        
        with ThreadPoolExecutor(max_workers=4) as executor:
            futures = {executor.submit(process_file, f): f for f in files}
            for future in as_completed(futures):
                f = futures[future]
                if verbose:
                    print(f"分析: {f.name}")
                try:
                    file_info, file_keywords, data_rows, file_type = future.result()
                    result['files'].append(file_info)
                    all_keywords.extend(file_keywords)
                    result['summary']['total_data_rows'] += data_rows
                    result['summary']['by_type'][file_type] = result['summary']['by_type'].get(file_type, 0) + 1
                except Exception as e:
                    logger.error(f"处理文件失败: {f}, 错误: {e}")
    else:
        # 串行处理
        for f in files:
            if verbose:
                print(f"分析: {f.name}")
            try:
                file_info, file_keywords, data_rows, file_type = process_file(f)
                result['files'].append(file_info)
                all_keywords.extend(file_keywords)
                result['summary']['total_data_rows'] += data_rows
                result['summary']['by_type'][file_type] = result['summary']['by_type'].get(file_type, 0) + 1
            except Exception as e:
                logger.error(f"处理文件失败: {f}, 错误: {e}")
    
    # 关键词统计
    keyword_counter = Counter(all_keywords)
    result['summary']['keywords'] = keyword_counter.most_common(30)
    
    # 工作类型推断
    result['summary']['work_type_guess'] = _guess_work_type(result['summary']['keywords'])
    
    if verbose:
        _print_analysis_report(result)
    
    return _make_response(True, data=result)


def _guess_work_type(keywords: List[Tuple[str, int]]) -> Dict[str, int]:
    """推断工作类型"""
    work_patterns = {
        '核查/审计': ['核查', '审计', '检查', '线索', '疑点', '违规', '问题', '风险'],
        '统计/汇总': ['汇总', '统计', '合计', '总计', '数量', '金额', '总数'],
        '报告/总结': ['报告', '总结', '情况', '分析', '建议', '附件', '说明'],
        '申报/审批': ['申报', '审批', '申请', '批准', '审核', '备案'],
        '考核/评估': ['考核', '评估', '评分', '得分', '排名', '绩效'],
        '台账/登记': ['台账', '登记', '记录', '清单', '名单', '花名册'],
        '模板/填报': ['模板', '表样', '填报', '格式', '样表'],
    }
    
    work_type_scores = {}
    for work_type, patterns in work_patterns.items():
        score = sum(freq for kw, freq in keywords for p in patterns if p in kw)
        if score > 0:
            work_type_scores[work_type] = score
    
    return work_type_scores


def _print_analysis_report(result: Dict):
    """打印分析报告"""
    print("\n" + "="*70)
    print("📊 目录分析报告")
    print("="*70)
    
    data = result.get('data', result)
    summary = data.get('summary', {})
    
    print(f"\n📁 目录: {data.get('directory', 'Unknown')}")
    print(f"📅 时间: {data.get('analyzed_at', 'Unknown')}")
    print(f"📄 文件: {summary.get('total_files', 0)} 个")
    print(f"   类型分布: {summary.get('by_type', {})}")
    print(f"📊 总数据行: {summary.get('total_data_rows', 0)}")
    
    if summary.get('work_type_guess'):
        print(f"\n💡 可能的工作类型:")
        for wt, score in sorted(summary['work_type_guess'].items(), key=lambda x: -x[1])[:3]:
            print(f"   • {wt} (匹配度: {score})")
    
    print(f"\n📄 文件详情:")
    for f in data.get('files', []):
        print(f"\n   {f['name']} ({f['type']}, {f['size']})")
        
        if f['type'] == 'excel' and 'sheets' in f.get('structure', {}):
            for sheet in f['structure']['sheets']:
                actual = sheet.get('actual_rows', sheet['rows'])
                print(f"      📊 {sheet['name']}: {actual}行×{sheet['cols']}列")
                if sheet.get('headers'):
                    print(f"         字段: {', '.join(sheet['headers'][:6])}...")
        
        elif f['type'] == 'word':
            struct = f.get('structure', {})
            print(f"      段落: {struct.get('paragraphs_count', 0)}, 表格: {struct.get('tables_count', 0)}")
        
        elif f['type'] == 'csv':
            struct = f.get('structure', {})
            print(f"      行数: {struct.get('rows', 0)}, 列数: {struct.get('cols', 0)}, 编码: {struct.get('encoding', 'unknown')}")
    
    if summary.get('keywords'):
        print(f"\n🔑 高频关键词: {', '.join(kw for kw, _ in summary['keywords'][:10])}")


# ============================================================================
# 缓存机制（优化 #18）
# ============================================================================

@lru_cache(maxsize=Config.CACHE_MAX_SIZE)
def _cached_file_info(filepath: str, mtime: float) -> Dict[str, Any]:
    """
    缓存文件信息（基于文件路径和修改时间）
    
    注意：mtime 参数用于缓存失效，当文件修改后缓存自动失效
    """
    return get_file_info(filepath)


def get_file_info_cached(filepath: str) -> Dict[str, Any]:
    """
    获取文件信息（带缓存）
    
    对同一文件的重复分析会使用缓存，文件修改后缓存自动失效
    """
    if not Config.CACHE_ENABLED:
        return get_file_info(filepath)
    
    filepath = Path(filepath)
    if not filepath.exists():
        return _make_response(False, error=f"文件不存在: {filepath}")
    
    mtime = filepath.stat().st_mtime
    return _cached_file_info(str(filepath.absolute()), mtime)


def clear_cache():
    """清除所有缓存"""
    _cached_file_info.cache_clear()
    logger.info("缓存已清除")


# ============================================================================
# 便捷函数：输出为JSON（供AI直接读取）
# ============================================================================

def to_json(result: Dict, indent: int = 2) -> str:
    """将结果转为JSON字符串，方便AI读取"""
    def default_serializer(obj):
        if isinstance(obj, datetime):
            return obj.isoformat()
        if isinstance(obj, Path):
            return str(obj)
        return str(obj)
    
    return json.dumps(result, ensure_ascii=False, indent=indent, default=default_serializer)


def save_analysis(result: Dict, filepath: str = None) -> str:
    """保存分析结果为JSON文件"""
    if filepath is None:
        filepath = f"analysis_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
    
    with open(filepath, 'w', encoding='utf-8') as f:
        f.write(to_json(result))
    
    logger.info(f"分析结果已保存: {filepath}")
    return filepath


# ============================================================================
# 导出功能
# ============================================================================

__all__ = [
    # 配置
    'Config',
    # 文件信息
    'get_file_info',
    'get_file_info_cached',
    'get_file_type',
    'scan_office_files',
    'convert_to_new_format',
    # Excel
    'read_excel',
    'write_excel',
    'analyze_excel_structure',
    'print_excel_analysis',
    # Word
    'read_word',
    'write_word',
    # PDF
    'read_pdf',
    # CSV（新增）
    'read_csv',
    'write_csv',
    # 目录分析
    'analyze_directory',
    # 分组统计
    'group_statistics',
    # 缓存
    'clear_cache',
    # 工具
    'to_json',
    'save_analysis',
]


# ============================================================================
# 分组统计（医保数据分析增强）
# ============================================================================

# 智能识别的分组列关键词
GROUP_COL_KEYWORDS = {
    '机构': ['医疗机构', '机构名称', '药店', '医院', '定点'],
    '人员': ['姓名', '参保人', '患者', '就诊人'],
    '药品': ['药品名称', '药品', '商品名'],
    '诊断': ['诊断', '疾病', '病种'],
    '时间': ['日期', '时间', '月份'],
}

# 智能识别的金额列关键词
AMOUNT_COL_KEYWORDS = ['金额', '费用', '支付', '报销', '自付', '统筹', '元']


def _find_group_column(headers: List[str], prefer_type: str = None) -> Optional[int]:
    """智能查找分组列"""
    if prefer_type and prefer_type in GROUP_COL_KEYWORDS:
        keywords = GROUP_COL_KEYWORDS[prefer_type]
        for i, h in enumerate(headers):
            if any(kw in h for kw in keywords):
                return i
    
    for col_type in ['机构', '人员', '药品', '诊断']:
        keywords = GROUP_COL_KEYWORDS[col_type]
        for i, h in enumerate(headers):
            if any(kw in h for kw in keywords):
                return i
    return None


def _find_amount_columns(headers: List[str]) -> List[int]:
    """智能查找金额列"""
    amount_cols = []
    for i, h in enumerate(headers):
        if any(kw in h for kw in AMOUNT_COL_KEYWORDS):
            amount_cols.append(i)
    return amount_cols


def _parse_number(val) -> Optional[float]:
    """解析数值"""
    if val is None:
        return None
    if isinstance(val, (int, float)):
        return float(val)
    if isinstance(val, str):
        cleaned = val.replace(',', '').replace('¥', '').replace('￥', '').strip()
        try:
            return float(cleaned)
        except:
            return None
    return None


def group_statistics(
    filepath: str,
    sheet_name: str = None,
    group_by: str = None,
    agg_cols: List[str] = None,
    top_n: int = 20
) -> Dict[str, Any]:
    """智能分组统计"""
    filepath = Path(filepath)
    
    accessible, error_msg = _check_file_accessible(filepath)
    if not accessible:
        return _make_response(False, error=error_msg)
    
    try:
        read_result = read_excel(str(filepath), sheet_name=sheet_name)
        if not read_result['success']:
            return read_result
        
        sheets = read_result['data']['sheets']
        if not sheets:
            return _make_response(False, error='没有找到数据')
        
        target_sheet = sheets[0]
        for s in sheets:
            if len(s.get('data', [])) > len(target_sheet.get('data', [])):
                target_sheet = s
        
        headers = target_sheet['headers']
        data = target_sheet['data']
        
        if not data:
            return _make_response(False, error='工作表没有数据')
        
        # 确定分组列
        group_col_idx = None
        group_col_name = None
        
        if group_by:
            for i, h in enumerate(headers):
                if group_by in h or h in group_by:
                    group_col_idx = i
                    group_col_name = h
                    break
            if group_col_idx is None:
                return _make_response(False, error=f'未找到分组列: {group_by}')
        else:
            group_col_idx = _find_group_column(headers)
            if group_col_idx is not None:
                group_col_name = headers[group_col_idx]
            else:
                return _make_response(False, error='无法自动识别分组列')
        
        # 确定聚合列
        agg_col_indices = []
        agg_col_names = []
        
        if agg_cols:
            for col_name in agg_cols:
                for i, h in enumerate(headers):
                    if col_name in h or h in col_name:
                        agg_col_indices.append(i)
                        agg_col_names.append(h)
                        break
        else:
            agg_col_indices = _find_amount_columns(headers)
            agg_col_names = [headers[i] for i in agg_col_indices]
        
        # 执行分组统计
        groups = {}
        
        for row in data:
            if group_col_idx >= len(row):
                continue
            
            group_val = row[group_col_idx]
            if group_val is None or (isinstance(group_val, str) and not group_val.strip()):
                group_val = '(空值)'
            else:
                group_val = str(group_val).strip()
            
            if group_val not in groups:
                groups[group_val] = {'count': 0, 'amounts': {col: [] for col in agg_col_names}}
            
            groups[group_val]['count'] += 1
            
            for idx, col_name in zip(agg_col_indices, agg_col_names):
                if idx < len(row):
                    num = _parse_number(row[idx])
                    if num is not None:
                        groups[group_val]['amounts'][col_name].append(num)
        
        # 计算汇总
        total_amounts = {col: 0.0 for col in agg_col_names}
        top_groups = []
        
        for group_name, group_data in groups.items():
            group_summary = {'name': group_name, 'count': group_data['count']}
            for col_name in agg_col_names:
                col_sum = sum(group_data['amounts'][col_name])
                group_summary[col_name] = round(col_sum, 2)
                total_amounts[col_name] += col_sum
            top_groups.append(group_summary)
        
        top_groups.sort(key=lambda x: x['count'], reverse=True)
        
        distribution = {'1-10条': 0, '11-50条': 0, '51-100条': 0, '101-500条': 0, '500条以上': 0}
        for g in top_groups:
            cnt = g['count']
            if cnt <= 10:
                distribution['1-10条'] += 1
            elif cnt <= 50:
                distribution['11-50条'] += 1
            elif cnt <= 100:
                distribution['51-100条'] += 1
            elif cnt <= 500:
                distribution['101-500条'] += 1
            else:
                distribution['500条以上'] += 1
        
        result = {
            'sheet_name': target_sheet['name'],
            'group_by': group_col_name,
            'total_rows': len(data),
            'total_groups': len(groups),
            'agg_columns': agg_col_names,
            'total_amounts': {k: round(v, 2) for k, v in total_amounts.items()},
            'top_groups': top_groups[:top_n],
            'distribution': distribution,
        }
        
        return _make_response(True, data=result)
        
    except Exception as e:
        logger.error(f"分组统计失败: {filepath}, 错误: {e}")
        return _make_response(False, error=str(e))


# ============================================================================
# 命令行入口
# ============================================================================

if __name__ == '__main__':
    import sys
    
    directory = sys.argv[1] if len(sys.argv) > 1 else '.'
    
    # 分析目录
    result = analyze_directory(directory, verbose=True)
    
    # 可选：保存为JSON
    if '--save' in sys.argv:
        filepath = save_analysis(result.get('data', result))
        print(f"\n💾 分析结果已保存: {filepath}")
